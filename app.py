import io
from datetime import date

import streamlit as st
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


st.set_page_config(page_title="Gestion des PV des marchés", layout="wide")


# =========================================================
# STATE
# =========================================================
def init_state():
    defaults = {
        "commune": "COMMUNE ASKAOUEN",
        "province": "PROVINCE DE TAROUDANT",
        "cercle": "CERCLE TALIOUINE",
        "caidat": "CAIDAT ASKAOUEN",
        "reference": "",
        "objet": "",
        "decision_no": "",
        "decision_date": date.today(),
        "session_date": date.today(),
        "session_time": "10:00",
        "session_place": "Salle de réunion de la Commune",
        "estimation": 0.0,
        "president": "",
        "publication_1": "",
        "publication_2": "",
        "portail_publication": "",
        "reprise_date": date.today(),
        "reprise_time": "10:00",
        "complement_limit_days": 7,
        "signature_place": "ASKAOUEN",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

    if "committee" not in st.session_state:
        st.session_state.committee = [
            {"name": "", "quality": "Président de la commune", "role": "PRESIDENT"},
            {"name": "", "quality": "Représentant du percepteur", "role": "MEMBRE"},
            {"name": "", "quality": "Directeur des services", "role": "MEMBRE"},
            {"name": "", "quality": "Technicien", "role": "MEMBRE"},
            {"name": "", "quality": "Service des marchés", "role": "MEMBRE"},
        ]

    if "subcommittee" not in st.session_state:
        st.session_state.subcommittee = [
            {"name": "", "quality": "Technicien à la commune"},
            {"name": "", "quality": "Technicien à la commune"},
            {"name": "", "quality": "Technicien à la commune"},
        ]

    if "bidders" not in st.session_state:
        st.session_state.bidders = [
            {
                "name": "",
                "admin_ok": True,
                "tech_ok": True,
                "financial_ok": True,
                "tech_score": 70.0,
                "amount": 0.0,
                "amount_rectified": 0.0,
                "excluded_reason_admin": "",
                "excluded_reason_tech": "",
                "complement_sent": "",
                "complement_received": "",
            }
        ]


# =========================================================
# HELPERS
# =========================================================
def fmt_date(d):
    return d.strftime("%d/%m/%Y") if isinstance(d, date) else str(d)


def committee_block():
    lines = []
    for m in st.session_state.committee:
        if m["name"].strip():
            lines.append(f"· {m['name']} : {m['quality']} -------------------------------------- {m['role']}")
    return "\n".join(lines) if lines else "· ..............................................................."


def bidder_names_numbered(filter_fn=None):
    items = []
    idx = 1
    for b in st.session_state.bidders:
        if b["name"].strip() and (filter_fn is None or filter_fn(b)):
            items.append(f"{idx}) {b['name']}")
            idx += 1
    return "\n".join(items) if items else "NEANT"


def admissible_names(kind="admin"):
    names = []
    reserved = []
    for b in st.session_state.bidders:
        if not b["name"].strip():
            continue
        if kind == "admin":
            if b["admin_ok"]:
                if b.get("excluded_reason_admin", "").strip():
                    reserved.append(b["name"])
                else:
                    names.append(b["name"])
        elif kind == "tech":
            if b["tech_ok"] and b["tech_score"] >= 70:
                if b.get("excluded_reason_tech", "").strip():
                    reserved.append(b["name"])
                else:
                    names.append(b["name"])
    return names, reserved


def excluded_rows(reason_key):
    rows = []
    for b in st.session_state.bidders:
        reason = b.get(reason_key, "").strip()
        if b["name"].strip() and reason:
            rows.append((b["name"], reason))
    return rows


def tech_score_rows():
    rows = []
    for b in st.session_state.bidders:
        if b["name"].strip() and b["admin_ok"]:
            rows.append((b["name"], str(b["tech_score"])))
    return rows


def financial_rows():
    rows = []
    for b in st.session_state.bidders:
        if b["name"].strip() and b["tech_ok"] and b["tech_score"] >= 70:
            rows.append((b["name"], f"{b['amount']:,.2f} DHS"))
    return rows


def rectification_rows():
    rows = []
    for b in st.session_state.bidders:
        if b["name"].strip() and b["tech_ok"] and b["tech_score"] >= 70:
            rectified = b["amount_rectified"] if b["amount_rectified"] > 0 else b["amount"]
            rows.append((b["name"], f"{b['amount']:,.2f} DHS", f"{rectified:,.2f} DHS"))
    return rows


def reference_price_data():
    amounts = []
    for b in st.session_state.bidders:
        if b["name"].strip() and b["tech_ok"] and b["tech_score"] >= 70:
            v = b["amount_rectified"] if b["amount_rectified"] > 0 else b["amount"]
            amounts.append((b["name"], v))
    if not amounts:
        return 0.0, [], None
    ref = (st.session_state.estimation + sum(v for _, v in amounts)) / (len(amounts) + 1)
    ranking = sorted(amounts, key=lambda x: abs(x[1] - ref))
    winner = ranking[0] if ranking else None
    return ref, amounts, winner


# =========================================================
# TEXT GENERATORS
# =========================================================
def build_pv1():
    adm_ok, adm_reserved = admissible_names("admin")
    tech_ok, tech_reserved = admissible_names("tech")
    sub = "\n".join(
        [f"· {m['name']} : {m['quality']}" for m in st.session_state.subcommittee if m["name"].strip()]
    ) or "· ................................ : technicien à la commune"

    return f"""ROYAUME DU MAROC

MINISTERE DE L’INTERIEUR

{st.session_state.province}
{st.session_state.cercle}
{st.session_state.caidat}
{st.session_state.commune}

PROCES VERBAL D'APPEL D'OFFRES OUVERT
SUR OFFRE DE PRIX N° : {st.session_state.reference}
1ère Séance Publique

Le {fmt_date(st.session_state.session_date)} à {st.session_state.session_time}, une commission d’appel d’offres, conformément à la décision de l’ordonnateur n° {st.session_state.decision_no} du {fmt_date(st.session_state.decision_date)}, est composée comme suit :

{committee_block()}

S’est réunie en séance publique dans {st.session_state.session_place}, en vue de procéder à l’ouverture des plis concernant l’appel d’offres ouvert sur offre de prix N° : {st.session_state.reference}, relatif à : {st.session_state.objet}.

Conformément à l’avis publié dans les journaux suivants :
· {st.session_state.publication_1 or '................................'}
· {st.session_state.publication_2 or '................................'}
· Mise en ligne au portail des marchés publics : {st.session_state.portail_publication or '................................'}

Le président cite les concurrents ayant déposé leurs plis :
{bidder_names_numbered()}

Le président s’assure de la présence des membres dont la présence est obligatoire.
Le président remet le support écrit contenant l’estimation des coûts détaillés des prestations dont le montant est fixé à {st.session_state.estimation:,.2f} DHS TTC.
Les membres de la commission paraphent le support de l’estimation des coûts des prestations.

Le président demande aux membres de la commission de formuler leurs réserves ou observations sur les vices éventuels qui entachent la procédure.
Le président ouvre les enveloppes extérieures des plis contenant les dossiers des concurrents, cite dans chacun la présence des enveloppes exigées, puis ouvre l’enveloppe portant la mention « dossiers administratif et technique ».

Cette formalité accomplie, la séance publique est suspendue, les concurrents et le public se retirent de la salle.

Ensuite, la commission se réunit à huis clos pour examiner les dossiers administratifs et techniques des concurrents.

Elle arrête ensuite la liste des concurrents admissibles :
A- Liste des concurrents admissibles sans réserves :
{chr(10).join([f'{i+1}) {x}' for i, x in enumerate(adm_ok)]) if adm_ok else 'NEANT'}
B- Liste des concurrents admissibles avec réserves :
{chr(10).join([f'{i+1}) {x}' for i, x in enumerate(adm_reserved)]) if adm_reserved else 'NEANT'}

La séance publique est alors reprise et le président donne lecture de la liste des soumissionnaires admissibles.
Le président procède ensuite à l’ouverture des enveloppes des soumissionnaires retenus portant la mention « offres techniques ».

Cette formalité accomplie, la séance publique est suspendue, les concurrents et le public se retirent de la salle.

Ensuite, la commission se réunit à huis clos pour examiner les offres techniques des concurrents.

Elle arrête ensuite la liste des concurrents admissibles :
A- Liste des concurrents admissibles sans réserves :
{chr(10).join([f'· {x}' for x in tech_ok]) if tech_ok else 'NEANT'}
B- Liste des concurrents admissibles avec réserves :
{chr(10).join([f'· {x}' for x in tech_reserved]) if tech_reserved else 'NEANT'}

Ensuite, et conformément aux dispositions de l’article 38 du décret n°2-22-431 du 08 mars 2023 relatif aux marchés publics, la commission a décidé de consulter une sous-commission technique pour examiner et analyser les offres techniques fournies par les concurrents.
La sous-commission technique est composée de :
{sub}

Le président de la commission suspend la séance et fixe la date de {fmt_date(st.session_state.reprise_date)} à {st.session_state.reprise_time} pour la reprise des travaux de la séance.
""".strip()


def build_rapport():
    below = [b["name"] for b in st.session_state.bidders if b["name"].strip() and b["tech_score"] < 70]
    above = [b["name"] for b in st.session_state.bidders if b["name"].strip() and b["tech_score"] >= 70]
    sub = "\n".join(
        [f"· {m['name']} : {m['quality']} ------------------------------------------------ MEMBRE"
         for m in st.session_state.subcommittee if m["name"].strip()]
    ) or "· ................................ : technicien à la commune ------------------------------------------------ MEMBRE"

    return f"""ROYAUME DU MAROC

MINISTERE DE L’INTERIEUR

{st.session_state.province}
{st.session_state.cercle}
{st.session_state.caidat}
{st.session_state.commune}

Rapport de la sous-commission technique
Appel d’offres ouvert {st.session_state.reference}

{st.session_state.objet}

EXAMEN DES OFFRES TECHNIQUES

Le {fmt_date(st.session_state.reprise_date)} à {st.session_state.reprise_time}, faisant suite à la séance d’ouverture des plis et à la décision du président de la commission, une sous-commission technique a été désignée conformément à l’article 38 du décret 2-22-431 relatif aux marchés publics.

Cette sous-commission technique est composée de :
{sub}

La liste des concurrents présentés par la commission d’ouverture pour l’examen des offres techniques est composée des concurrents suivants :
{bidder_names_numbered(lambda b: b['admin_ok'])}

Conclusion :
Après l’examen des offres techniques des concurrents :
La sous-commission technique arrête la liste des concurrents dont la note des offres techniques est supérieure ou égale à 70 points :
{chr(10).join([f'· {x}' for x in above]) if above else 'NEANT'}

La sous-commission technique arrête la liste des concurrents dont la note des offres techniques est inférieure à 70 points :
{chr(10).join([f'· {x}' for x in below]) if below else 'NEANT'}
""".strip()


def build_pv2():
    ref, amounts, winner = reference_price_data()
    ranking_lines = [
        f"{i}. Offre financière {name} = {amount:,.2f} DHS"
        for i, (name, amount) in enumerate(sorted(amounts, key=lambda x: abs(x[1] - ref)), start=1)
    ]
    winner_text = f"{winner[0]} = {winner[1]:,.2f} DHS" if winner else "................................"

    return f"""ROYAUME DU MAROC

MINISTERE DE L’INTERIEUR

{st.session_state.province}
{st.session_state.cercle}
{st.session_state.caidat}
{st.session_state.commune}

PROCES VERBAL D'APPEL D'OFFRES OUVERT
SUR OFFRE DE PRIX N° : {st.session_state.reference}
2ème Séance Publique

Conformément à la décision de l’ordonnateur n° {st.session_state.decision_no} du {fmt_date(st.session_state.decision_date)}, le {fmt_date(st.session_state.reprise_date)} à {st.session_state.reprise_time}, la commission composée comme suit :
{committee_block()}

s’est réunie en séance publique à {st.session_state.session_place} en vue d’étudier le rapport de la sous-commission technique.

Conformément aux critères d’évaluation des offres, les concurrents ayant obtenu une note inférieure à 70 points seront écartés.

Elle procède au calcul du prix de référence comme suit :
· Estimation = {st.session_state.estimation:,.2f} DHS
{chr(10).join([f'· Offre financière {n} = {v:,.2f} DHS' for n, v in amounts]) if amounts else '· Aucune offre admissible'}
· Le prix de référence = {ref:,.2f} DHS

La commission procède au classement des offres des concurrents au regard du prix de référence :
{chr(10).join(ranking_lines) if ranking_lines else 'NEANT'}

L’offre économiquement la plus avantageuse à proposer au maître d’ouvrage est celle qui est la plus proche par défaut du prix de référence, à savoir :
{winner_text}

La commission invite, par voie électronique, le concurrent ayant présenté l’offre économiquement la plus avantageuse, dans un délai de {st.session_state.complement_limit_days} jours après réception de la lettre, à produire le complément du dossier administratif.
""".strip()


def build_pv3():
    _, _, winner = reference_price_data()
    winner_name = winner[0] if winner else "................................"
    winner_amount = f"{winner[1]:,.2f} DHS" if winner else "................................"
    invited = next((b for b in st.session_state.bidders if b["name"] == winner_name), None)
    sent = invited["complement_sent"] if invited else ""
    received = invited["complement_received"] if invited else ""

    return f"""ROYAUME DU MAROC

MINISTERE DE L’INTERIEUR

{st.session_state.province}
{st.session_state.cercle}
{st.session_state.caidat}
{st.session_state.commune}

PROCES VERBAL D'APPEL D'OFFRES OUVERT
SUR OFFRE DE PRIX N° : {st.session_state.reference}
3ème Séance Publique

Le {fmt_date(st.session_state.session_date)} à {st.session_state.session_time}, une commission d’appel d’offres, conformément à la décision de l’ordonnateur n° {st.session_state.decision_no} du {fmt_date(st.session_state.decision_date)}, et composée comme suit :
{committee_block()}

s’est réunie en séance publique à {st.session_state.session_place}, en vue de procéder à l’ouverture du complément du dossier administratif de l’attributaire de l’appel d’offres ouvert sur offre de prix N° : {st.session_state.reference}, relatif à : {st.session_state.objet}.

La commission s’assure du support ayant servi de moyen d’invitation du concurrent concerné {winner_name}.
Date d’envoi de la lettre : {sent or '................................'}
Elle vérifie les pièces et la réponse reçue :
Dossier déposé le : {received or '................................'}

La commission examine les pièces complémentaires du dossier administratif et la réponse reçue, les juge acceptables, et décide de proposer au maître d’ouvrage de retenir l’offre du concurrent ayant présenté l’offre la plus avantageuse, à savoir :
{winner_name} pour un montant de {winner_amount}
""".strip()


def build_os_notification(company_name):
    return f"""ROYAUME DU MAROC

MINISTERE DE L'INTERIEUR
{st.session_state.province}
{st.session_state.cercle}
{st.session_state.caidat}
{st.session_state.commune}

ORDRE DE SERVICE DE LA NOTIFICATION
DE L’APPROBATION DU MARCHE N° : {st.session_state.reference}

Le maître d’ouvrage représenté par {st.session_state.president},
informe l’entreprise {company_name} que le marché ayant pour objet : {st.session_state.objet}
est approuvé à la date du : ................................
Par conséquent, l’intéressé est invité à acquitter les droits de timbre dus au titre du présent marché, conformément à la législation en vigueur.
""".strip()


def build_os_commencement(company_name):
    return f"""ROYAUME DU MAROC

MINISTERE DE L'INTERIEUR
{st.session_state.province}
{st.session_state.cercle}
{st.session_state.caidat}
{st.session_state.commune}

ORDRE DE SERVICE A L’ENTREPRENEUR POUR COMMENCEMENT DES TRAVAUX
N° : {st.session_state.reference}

Le maître d’ouvrage représenté par {st.session_state.president},
informe l’entreprise {company_name} que le marché ayant pour objet : {st.session_state.objet} est approuvé.

Par conséquent, l’intéressé est invité à commencer les travaux objet du présent marché à compter du : ............................
""".strip()


# =========================================================
# DOCX HELPERS
# =========================================================
def set_cell_text(cell, text, bold=False, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")
        borders.append(elem)
    tblPr.append(borders)


def add_header_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "ROYAUME DU MAROC",
        "MINISTERE DE L’INTERIEUR",
        st.session_state.province,
        st.session_state.cercle,
        st.session_state.caidat,
        st.session_state.commune,
    ]:
        r = p.add_run(line + "\n")
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(11)


def add_title(doc, title, subtitle=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(13)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.bold = True
        r2.font.name = "Arial"
        r2.font.size = Pt(11)


def add_paragraph(doc, text, bold=False, center=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(11)
    return p


def add_simple_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, center=True)

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, center=(i > 0))
    return table


def add_signature_block(doc):
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"Fait à {st.session_state.signature_place}, le : {fmt_date(st.session_state.session_date)}")
    r.font.name = "Arial"
    r.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("SIGNE : LE PRESIDENT")
    r2.bold = True
    r2.font.name = "Arial"
    r2.font.size = Pt(11)

    doc.add_paragraph("")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("LES MEMBRES")
    r3.bold = True
    r3.font.name = "Arial"
    r3.font.size = Pt(11)


def create_base_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    return doc


# =========================================================
# DOCX BUILDERS
# =========================================================
def build_docx_pv1():
    doc = create_base_doc()
    add_header_block(doc)
    add_title(doc, "PROCES VERBAL D'APPEL D'OFFRES OUVERT", f"SUR OFFRE DE PRIX N° : {st.session_state.reference}")
    add_title(doc, "1ère Séance Publique")

    add_paragraph(doc, f"Le {fmt_date(st.session_state.session_date)} à {st.session_state.session_time}, une commission d’appel d’offres, conformément à la décision de l’ordonnateur n° {st.session_state.decision_no} du {fmt_date(st.session_state.decision_date)}, est composée comme suit :")

    members = [(m["name"], m["quality"], m["role"]) for m in st.session_state.committee if m["name"].strip()]
    add_simple_table(doc, ["Nom", "Qualité", "Rôle"], members or [("...", "...", "...")])

    add_paragraph(doc, f"S’est réunie en séance publique dans {st.session_state.session_place}, en vue de procéder à l’ouverture des plis concernant l’appel d’offres ouvert sur offre de prix N° : {st.session_state.reference}, relatif à : {st.session_state.objet}.")
    add_paragraph(doc, "Conformément à l’avis publié dans les journaux suivants :")
    pubs = [
        (st.session_state.publication_1 or "...",),
        (st.session_state.publication_2 or "...",),
        (f"Portail des marchés publics : {st.session_state.portail_publication or '...'}",),
    ]
    add_simple_table(doc, ["Publications"], pubs)

    add_paragraph(doc, "Le président cite les concurrents ayant déposé leurs plis :")
    bidder_rows = [(b["name"],) for b in st.session_state.bidders if b["name"].strip()]
    add_simple_table(doc, ["Concurrents"], bidder_rows or [("NEANT",)])

    excl_admin = excluded_rows("excluded_reason_admin")
    add_paragraph(doc, "Concurrents écartés après examen administratif et technique :")
    add_simple_table(doc, ["Concurrents", "Motif d’écartement"], excl_admin or [("NEANT", "NEANT")])

    adm_ok, adm_reserved = admissible_names("admin")
    add_paragraph(doc, "Liste des concurrents admissibles sans réserves :")
    add_simple_table(doc, ["Concurrents"], [(x,) for x in adm_ok] or [("NEANT",)])
    add_paragraph(doc, "Liste des concurrents admissibles avec réserves :")
    add_simple_table(doc, ["Concurrents"], [(x,) for x in adm_reserved] or [("NEANT",)])

    excl_tech = excluded_rows("excluded_reason_tech")
    add_paragraph(doc, "Concurrents écartés après examen des offres techniques :")
    add_simple_table(doc, ["Concurrents", "Motif d’écartement"], excl_tech or [("NEANT", "NEANT")])

    tech_ok, tech_reserved = admissible_names("tech")
    add_paragraph(doc, "Liste des concurrents admissibles sans réserves :")
    add_simple_table(doc, ["Concurrents"], [(x,) for x in tech_ok] or [("NEANT",)])
    add_paragraph(doc, "Liste des concurrents admissibles avec réserves :")
    add_simple_table(doc, ["Concurrents"], [(x,) for x in tech_reserved] or [("NEANT",)])

    add_paragraph(doc, "Sous-commission technique :")
    sub_rows = [(m["name"], m["quality"]) for m in st.session_state.subcommittee if m["name"].strip()]
    add_simple_table(doc, ["Nom", "Qualité"], sub_rows or [("...", "Technicien à la commune")])

    add_paragraph(doc, f"Le président de la commission suspend la séance et fixe la date de {fmt_date(st.session_state.reprise_date)} à {st.session_state.reprise_time} pour la reprise des travaux de la séance.")
    add_signature_block(doc)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def build_docx_rapport():
    doc = create_base_doc()
    add_header_block(doc)
    add_title(doc, "Rapport de la sous-commission technique", f"Appel d’offres ouvert {st.session_state.reference}")
    add_paragraph(doc, st.session_state.objet, bold=True, center=True)
    add_paragraph(doc, "EXAMEN DES OFFRES TECHNIQUES", bold=True, center=True)

    add_paragraph(doc, f"Le {fmt_date(st.session_state.reprise_date)} à {st.session_state.reprise_time}, faisant suite à la séance d’ouverture des plis et à la décision du président de la commission, une sous-commission technique a été désignée conformément à l’article 38 du décret 2-22-431 relatif aux marchés publics.")

    sub_rows = [(m["name"], m["quality"]) for m in st.session_state.subcommittee if m["name"].strip()]
    add_simple_table(doc, ["Nom", "Qualité"], sub_rows or [("...", "Technicien à la commune")])

    rows = [(b["name"],) for b in st.session_state.bidders if b["name"].strip() and b["admin_ok"]]
    add_paragraph(doc, "Concurrents présentés pour l’examen des offres techniques :")
    add_simple_table(doc, ["Concurrents"], rows or [("NEANT",)])

    score_rows = tech_score_rows()
    add_paragraph(doc, "Fiche de notation et évaluation :")
    add_simple_table(doc, ["Concurrents", "Note technique"], score_rows or [("NEANT", "0")])

    add_signature_block(doc)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def build_docx_pv2():
    doc = create_base_doc()
    add_header_block(doc)
    add_title(doc, "PROCES VERBAL D'APPEL D'OFFRES OUVERT", f"SUR OFFRE DE PRIX N° : {st.session_state.reference}")
    add_title(doc, "2ème Séance Publique")

    ref, amounts, winner = reference_price_data()
    winner_text = f"{winner[0]} - {winner[1]:,.2f} DHS" if winner else "NEANT"

    add_paragraph(doc, f"Conformément à la décision de l’ordonnateur n° {st.session_state.decision_no} du {fmt_date(st.session_state.decision_date)}, le {fmt_date(st.session_state.reprise_date)} à {st.session_state.reprise_time}, la commission s’est réunie en séance publique à {st.session_state.session_place} en vue d’étudier le rapport de la sous-commission technique.")

    add_paragraph(doc, "Liste et notes des offres techniques des concurrents admissibles :")
    add_simple_table(doc, ["Concurrents", "Note Technique"], tech_score_rows() or [("NEANT", "0")])

    add_paragraph(doc, "Montants des actes d’engagement :")
    add_simple_table(doc, ["Concurrents", "Montant"], financial_rows() or [("NEANT", "0 DHS")])

    add_paragraph(doc, "Rectification des montants :")
    add_simple_table(doc, ["Concurrents", "Avant rectification", "Montant rectifié"], rectification_rows() or [("NEANT", "0 DHS", "0 DHS")])

    calc_rows = [("Estimation", f"{st.session_state.estimation:,.2f} DHS")]
    calc_rows.extend([(n, f"{v:,.2f} DHS") for n, v in amounts] or [("Aucune offre admissible", "0 DHS")])
    calc_rows.append(("Prix de référence", f"{ref:,.2f} DHS"))
    add_paragraph(doc, "Calcul du prix de référence :")
    add_simple_table(doc, ["Libellé", "Montant"], calc_rows)

    ranking = sorted(amounts, key=lambda x: abs(x[1] - ref))
    add_paragraph(doc, "Classement des offres :")
    add_simple_table(doc, ["Rang", "Concurrent", "Montant"], [(i, n, f"{v:,.2f} DHS") for i, (n, v) in enumerate(ranking, start=1)] or [("1", "NEANT", "0 DHS")])

    add_paragraph(doc, f"L’offre économiquement la plus avantageuse est : {winner_text}", bold=True)
    add_signature_block(doc)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def build_docx_pv3():
    doc = create_base_doc()
    add_header_block(doc)
    add_title(doc, "PROCES VERBAL D'APPEL D'OFFRES OUVERT", f"SUR OFFRE DE PRIX N° : {st.session_state.reference}")
    add_title(doc, "3ème Séance Publique")

    _, _, winner = reference_price_data()
    winner_name = winner[0] if winner else "NEANT"
    winner_amount = f"{winner[1]:,.2f} DHS" if winner else "0 DHS"
    invited = next((b for b in st.session_state.bidders if b["name"] == winner_name), None)
    sent = invited["complement_sent"] if invited else ""
    received = invited["complement_received"] if invited else ""

    add_paragraph(doc, f"Le {fmt_date(st.session_state.session_date)} à {st.session_state.session_time}, la commission s’est réunie en séance publique à {st.session_state.session_place} pour l’ouverture du complément du dossier administratif de l’attributaire relatif à : {st.session_state.objet}.")
    add_simple_table(doc, ["Concurrent concerné", "Date d’envoi", "Date de dépôt"], [(winner_name, sent or "...", received or "...")])
    add_paragraph(doc, f"La commission décide de proposer au maître d’ouvrage de retenir l’offre du concurrent : {winner_name} pour un montant de {winner_amount}.", bold=True)

    add_signature_block(doc)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def build_docx_os_notification(company_name):
    doc = create_base_doc()
    add_header_block(doc)
    add_title(doc, "ORDRE DE SERVICE DE LA NOTIFICATION", f"DE L’APPROBATION DU MARCHE N° : {st.session_state.reference}")
    add_paragraph(doc, f"Le maître d’ouvrage représenté par {st.session_state.president}, informe l’entreprise {company_name} que le marché ayant pour objet : {st.session_state.objet} est approuvé à la date du : ................................")
    add_paragraph(doc, "Par conséquent, l’intéressé est invité à acquitter les droits de timbre dus au titre du présent marché, conformément à la législation en vigueur.")
    add_signature_block(doc)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def build_docx_os_commencement(company_name):
    doc = create_base_doc()
    add_header_block(doc)
    add_title(doc, "ORDRE DE SERVICE A L’ENTREPRENEUR POUR COMMENCEMENT DES TRAVAUX", f"N° : {st.session_state.reference}")
    add_paragraph(doc, f"Le maître d’ouvrage représenté par {st.session_state.president}, informe l’entreprise {company_name} que le marché ayant pour objet : {st.session_state.objet} est approuvé.")
    add_paragraph(doc, "Par conséquent, l’intéressé est invité à commencer les travaux objet du présent marché à compter du : ............................")
    add_signature_block(doc)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


# =========================================================
# UI
# =========================================================
init_state()
st.title("📄 Générateur des PV des marchés publics")

tab1, tab2, tab3, tab4 = st.tabs(["Données générales", "Commission", "Concurrents", "Génération"])

with tab1:
    c1, c2 = st.columns(2)

    with c1:
        st.session_state.commune = st.text_input("Commune", st.session_state.commune)
        st.session_state.province = st.text_input("Province", st.session_state.province)
        st.session_state.cercle = st.text_input("Cercle", st.session_state.cercle)
        st.session_state.caidat = st.text_input("Caidat", st.session_state.caidat)
        st.session_state.reference = st.text_input("Référence AO / marché", st.session_state.reference)
        st.session_state.objet = st.text_area("Objet", st.session_state.objet, height=100)
        st.session_state.estimation = st.number_input("Estimation TTC (DHS)", min_value=0.0, value=float(st.session_state.estimation), step=1000.0)

    with c2:
        st.session_state.decision_no = st.text_input("N° décision ordonnateur", st.session_state.decision_no)
        st.session_state.decision_date = st.date_input(
            "Date décision",
            value=st.session_state.decision_date,
            key="decision_date_widget"
        )
        st.session_state.session_date = st.date_input(
            "Date séance",
            value=st.session_state.session_date,
            key="session_date_widget"
        )
        st.session_state.session_time = st.text_input("Heure séance", st.session_state.session_time)
        st.session_state.session_place = st.text_input("Lieu séance", st.session_state.session_place)
        st.session_state.president = st.text_input("Président maître d’ouvrage", st.session_state.president)
        st.session_state.publication_1 = st.text_input("Publication journal 1", st.session_state.publication_1)
        st.session_state.publication_2 = st.text_input("Publication journal 2", st.session_state.publication_2)
        st.session_state.portail_publication = st.text_input("Publication portail marchés publics", st.session_state.portail_publication)
        st.session_state.reprise_date = st.date_input(
            "Date reprise",
            value=st.session_state.reprise_date,
            key="reprise_date_widget"
        )
        st.session_state.reprise_time = st.text_input("Heure reprise", st.session_state.reprise_time)
        st.session_state.signature_place = st.text_input("Lieu de signature", st.session_state.signature_place)

with tab2:
    st.subheader("Commission d’appel d’offres")
    committee_n = st.number_input("Nombre de membres", min_value=1, max_value=10, value=len(st.session_state.committee), step=1)
    while len(st.session_state.committee) < committee_n:
        st.session_state.committee.append({"name": "", "quality": "", "role": "MEMBRE"})
    while len(st.session_state.committee) > committee_n:
        st.session_state.committee.pop()

    for i, m in enumerate(st.session_state.committee):
        a, b, c = st.columns([2, 2, 1])
        m["name"] = a.text_input(f"Nom membre {i+1}", m["name"], key=f"cm_name_{i}")
        m["quality"] = b.text_input(f"Qualité {i+1}", m["quality"], key=f"cm_quality_{i}")
        m["role"] = c.selectbox(f"Rôle {i+1}", ["PRESIDENT", "MEMBRE"], index=0 if m["role"] == "PRESIDENT" else 1, key=f"cm_role_{i}")

    st.subheader("Sous-commission technique")
    sub_n = st.number_input("Nombre membres sous-commission", min_value=1, max_value=5, value=len(st.session_state.subcommittee), step=1)
    while len(st.session_state.subcommittee) < sub_n:
        st.session_state.subcommittee.append({"name": "", "quality": "Technicien à la commune"})
    while len(st.session_state.subcommittee) > sub_n:
        st.session_state.subcommittee.pop()

    for i, m in enumerate(st.session_state.subcommittee):
        a, b = st.columns(2)
        m["name"] = a.text_input(f"Nom sous-commission {i+1}", m["name"], key=f"scm_name_{i}")
        m["quality"] = b.text_input(f"Qualité sous-commission {i+1}", m["quality"], key=f"scm_quality_{i}")

with tab3:
    st.subheader("Concurrents")
    bid_n = st.number_input("Nombre de concurrents", min_value=1, max_value=20, value=len(st.session_state.bidders), step=1)
    while len(st.session_state.bidders) < bid_n:
        st.session_state.bidders.append({
            "name": "",
            "admin_ok": True,
            "tech_ok": True,
            "financial_ok": True,
            "tech_score": 70.0,
            "amount": 0.0,
            "amount_rectified": 0.0,
            "excluded_reason_admin": "",
            "excluded_reason_tech": "",
            "complement_sent": "",
            "complement_received": "",
        })
    while len(st.session_state.bidders) > bid_n:
        st.session_state.bidders.pop()

    for i, b in enumerate(st.session_state.bidders):
        st.markdown(f"### Concurrent {i+1}")
        c1, c2 = st.columns(2)
        b["name"] = c1.text_input("Nom", b["name"], key=f"bid_name_{i}")
        b["tech_score"] = c2.number_input("Note technique", min_value=0.0, max_value=100.0, value=float(b["tech_score"]), step=1.0, key=f"bid_score_{i}")

        c3, c4 = st.columns(2)
        b["admin_ok"] = c3.checkbox("Admis administratif", value=b["admin_ok"], key=f"bid_admin_{i}")
        b["tech_ok"] = c4.checkbox("Admis technique", value=b["tech_ok"], key=f"bid_tech_{i}")

        c5, c6 = st.columns(2)
        b["amount"] = c5.number_input("Montant engagement", min_value=0.0, value=float(b["amount"]), step=1000.0, key=f"bid_amount_{i}")
        b["amount_rectified"] = c6.number_input("Montant rectifié", min_value=0.0, value=float(b["amount_rectified"]), step=1000.0, key=f"bid_amount_rect_{i}")

        c7, c8 = st.columns(2)
        b["excluded_reason_admin"] = c7.text_input("Motif écartement administratif", b["excluded_reason_admin"], key=f"bid_exc_admin_{i}")
        b["excluded_reason_tech"] = c8.text_input("Motif écartement technique", b["excluded_reason_tech"], key=f"bid_exc_tech_{i}")

        c9, c10 = st.columns(2)
        b["complement_sent"] = c9.text_input("Date envoi invitation complément", b["complement_sent"], key=f"bid_sent_{i}")
        b["complement_received"] = c10.text_input("Date dépôt complément", b["complement_received"], key=f"bid_received_{i}")

with tab4:
    st.subheader("Génération des documents")
    ref, _, winner = reference_price_data()
    if winner:
        st.info(f"Offre la plus proche du prix de référence ({ref:,.2f} DHS) : {winner[0]} - {winner[1]:,.2f} DHS")

    doc_type = st.selectbox(
        "Choisir le document",
        ["PV 1ère séance", "Rapport sous-commission", "PV 2ème séance", "PV 3ème séance", "OS notification", "OS commencement"]
    )

    winner_name = winner[0] if winner else ""
    company_name = st.text_input("Entreprise concernée", value=winner_name)

    if st.button("Générer"):
        text_generators = {
            "PV 1ère séance": build_pv1,
            "Rapport sous-commission": build_rapport,
            "PV 2ème séance": build_pv2,
            "PV 3ème séance": build_pv3,
            "OS notification": lambda: build_os_notification(company_name),
            "OS commencement": lambda: build_os_commencement(company_name),
        }

        docx_generators = {
            "PV 1ère séance": build_docx_pv1,
            "Rapport sous-commission": build_docx_rapport,
            "PV 2ème séance": build_docx_pv2,
            "PV 3ème séance": build_docx_pv3,
            "OS notification": lambda: build_docx_os_notification(company_name),
            "OS commencement": lambda: build_docx_os_commencement(company_name),
        }

        text = text_generators[doc_type]()
        docx_bytes = docx_generators[doc_type]()

        st.text_area("Document généré", text, height=450)

        safe_name = doc_type.lower().replace(" ", "_").replace("è", "e").replace("é", "e")

        st.download_button(
            "Télécharger TXT",
            text,
            file_name=f"{safe_name}.txt",
            mime="text/plain"
        )

        st.download_button(
            "Télécharger Word (.docx)",
            docx_bytes,
            file_name=f"{safe_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
